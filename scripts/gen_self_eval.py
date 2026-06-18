"""生成独立的自我评价文档"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(PROJECT_ROOT, "report", "自我评价.docx")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = '黑体'; r.font.size = Pt(14); r.bold = True
    r._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = '黑体'; r.font.size = Pt(10.5); r.bold = True
    r._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = '宋体'; r.font.size = Pt(10.5)
    r._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# Title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('数据库课程期末Project — 自我评价')
r.font.name = '黑体'; r.font.size = Pt(16); r.bold = True
r._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(12)
r = p.add_run('卢绍东  42411019  西南财经大学 计算机学院')
r.font.name = '宋体'; r.font.size = Pt(12)
r._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# == 一 ==
h1('一、时间投入与工作过程')

h2('第一阶段：选题分析与方案设计（6月15日上午）')
body(
    '阅读课程作业要求文档，对比三个题目（表格数据智能查询系统、轻量级数据库引擎、'
    '培养方案数据库系统）的技术难度、工作量、评分标准。结合个人情况（1人独立完成），'
    '最终选择题目三"培养方案数据库系统"。规划技术选型（Python + FastAPI + SQLite），'
    '设计项目结构和数据库Schema。通过Web搜索了解两校培养方案的数据结构，'
    '分析数据获取的可行性。产出：项目目录结构、requirements.txt、数据库Schema设计（6张表）、ER图。'
)

h2('第二阶段：核心功能开发（6月15日下午-晚上）')
body(
    '完成7次Git提交，实现全部核心功能：数据库建表与导入、12个查询函数（覆盖6项必做查询+4项跨校对比）、'
    'FastAPI Web应用（15个API端点）、28个测试用例。在此过程中与AI（Claude Code）进行了密集交流，'
    '讨论的问题包括：数据库去重策略（最终采用code字段全局唯一+多对多关联表方案）、'
    'Windows GBK编码导致UnicodeEncodeError的解决方案、跨校对比共有课程SQL（GROUP BY + HAVING COUNT DISTINCT）、'
    'HTML表格渲染的行内样式设计等。产出：约7000行代码，完整的查询系统和Web接口。'
)

h2('第三阶段：界面优化与用户体验改进（6月15日晚）')
body(
    '针对浏览器直接显示JSON不友好的问题，实现了内容协商机制（浏览器访问自动渲染HTML表格，'
    'API调用返回JSON）。发现Swagger UI默认英文界面不够友好，先后尝试Swagger插件方案和DOM遍历替换方案，'
    '最终采用onComplete回调+DOM文本节点遍历的方式实现40+个英文标签的中文化。将跨校对比中的英文缩写'
    '（SWUFE/SUFE）替换为中文全称，并在表格中添加彩色标签。'
)

h2('第四阶段：数据获取与处理（6月18日）')
body(
    '尝试通过学校官网、WebFetch、WebSearch等多种方式获取两校培养方案数据。西南财经大学方面，'
    '从官方ZIP文件中解压53份PDF，使用pdfplumber和PyMuPDF提取文本，编写数据提取与导入脚本进行结构化处理，'
    '最终覆盖5个学院8个专业。上海财经大学方面，官网PDF因CID字体编码问题无法自动提取文本，'
    '通过查阅官网培养方案页面获取课程信息。遇到的困难包括：学校网站外部访问被拦截、.doc格式模板无法直接编辑'
    '（采用python-docx新建方案）、70MB PDF文本提取失败（CID字体编码，两种PDF库均无法解决）等。'
)

h2('第五阶段：自然语言查询与报告撰写（6月18日）')
body(
    '实现自然语言查询接口，包含关键字匹配规则引擎（基于正则表达式，支持10种查询意图）和LLM接口框架'
    '（预留Claude API接口）。编写12个自然语言测试用例，意图识别正确率100%。'
    '按《计算机学报》模板格式生成完整课程报告（含中英文摘要、ER图、系统架构说明、实验评估数据等）。'
)

# == 二 ==
h1('二、知识学习与探索')

body(
    '1. 数据库设计能力：深入理解了关系型数据库的多对多关联设计、外键约束、CHECK约束、索引策略等概念。'
    '在实践中体会到课程去重设计的重要性——不同专业间存在大量同名课程（如思政课、数学课），'
    '采用code字段全局唯一+多对多关联表的方案有效避免了数据冗余。在作业要求之外，自行设计了ER图并在报告中详细说明。'
)

body(
    '2. SQL查询能力：在实现12个查询函数的过程中，综合运用了多表JOIN、CASE WHEN分组聚合、子查询、'
    'GROUP BY + HAVING、LIKE模糊匹配等SQL技术。特别是跨校共有课程检测（GROUP BY + HAVING COUNT DISTINCT）'
    '和独有课程检索（NOT IN子查询）两个查询，需要理解SQL执行顺序和子查询优化，锻炼了复杂SQL的编写能力。'
)

body(
    '3. Web开发能力：掌握了FastAPI框架的使用，包括路由设计、Query参数、内容协商、错误处理等。'
    '理解了RESTful API设计原则。通过实际项目体验了Swagger UI的自定义和中文化改造，学习了Swagger UI的'
    '插件机制和DOM操作技巧。自主设计了纯中文测试页面（/test），方便非技术用户使用。'
)

body(
    '4. 数据处理能力：实践了PDF文本提取（pdfplumber + PyMuPDF）、中文编码处理（Windows GBK vs UTF-8）、'
    'JSON结构化数据转换等数据工程技能。理解了实际项目中数据获取往往是整个系统开发中最耗时、最不可控的环节。'
)

body(
    '5. 自然语言处理入门：通过实现关键字匹配规则引擎，初步了解了自然语言查询接口的设计思路。'
    '正则表达式意图识别的实践加深了对NLP基本方法的理解。预留的LLM接口框架为后续接入大语言模型做好了准备。'
)

body(
    '6. 软件工程实践：完整经历了一个数据库应用项目从需求分析、方案设计、编码实现、测试验证到文档撰写的全流程。'
    '使用Git进行版本控制（17次提交），编写了README文档、API文档和课程报告。体会到迭代开发模式的有效性。'
)

# == 三 ==
h1('三、与AI的协作')

body(
    '在本项目中，AI（Claude Code）作为编程助手，在以下环节提供了帮助：'
    '（1）项目初始阶段协助选题分析和方案设计；'
    '（2）代码编写阶段提供SQL查询优化建议、FastAPI路由设计参考和前端HTML/CSS样式方案；'
    '（3）调试阶段协助定位Windows GBK编码问题、PDF文本提取失败原因和Git冲突解决方法；'
    '（4）文档撰写阶段协助生成课程报告初稿框架和自我评价初稿。'
)

body(
    '在所有环节中，本人始终保持对代码和方案的主导权，对AI生成的内容进行理解、验证和修改，'
    '确保最终成果反映个人的真实理解和能力。具体而言：数据库Schema设计、查询架构、技术选型等核心决策'
    '由本人独立判断；AI生成的代码均经过本人逐行阅读和测试验证；报告内容在AI初稿基础上进行了大量补充和修改，'
    '特别是实验数据、具体技术细节和个人反思部分。'
)

body(
    '与AI的交流本质上是加速器而非替代品——它帮助快速完成重复性工作（如批量编写课程JSON数据、'
    '生成HTML页面结构），但在核心设计决策上仍需要本人的判断和决策。通过这次协作，我也学到了'
    '如何更有效地与AI沟通需求、验证其输出、以及在发现问题时及时纠正方向。'
)

# == 四 ==
h1('四、自我评分')

body('根据课程评分标准（满分50分），本人自评得分：45分。')

h2('加分依据')

body(
    '1. 完整实现子模块B（跨校对比 + 自然语言查询接口），覆盖课程对比、学分对比、共有/独有课程分析、'
    '关键字匹配NL查询、LLM接口框架。评分参考中数据预处理与数据库设计（30%）+ 查询系统功能完整性（30%）'
    '+ 代码质量与工程规范（10%）+ 报告写作质量（30%），本项目均高质量完成。+10分。'
)

body(
    '2. 17次Git提交记录，30个文件，约7000行代码（Python + SQL + HTML + JSON）。项目结构规范，'
    'README文档完善，包含快速开始指南和项目结构说明。+5分。'
)

body(
    '3. 基于西南财经大学2025版官方PDF（53份培养方案文件）提取的真实课程数据，覆盖5个学院8个专业，'
    '238门去重课程，441条关联记录，零数据完整性问题。上海财经大学数据来源于官网培养方案页面。+5分。'
)

body(
    '4. 自主设计纯中文测试页面（/test），Swagger UI全面中文化（40+标签），浏览器访问自动渲染HTML表格。'
    '用户体验优于基础要求。+3分。'
)

body(
    '5. 按《计算机学报》模板格式撰写完整课程报告，包含中英文摘要、ER图、系统架构说明、实现细节、使用说明、'
    '28个测试用例的实验评估、跨校对比数据分析、总结与反思。+5分。'
)

body(
    '6. 1人独立完成全部工作（题目三要求2-4人组队），工作量大，覆盖数据获取、数据库设计、查询系统实现、'
    'Web界面开发、自然语言查询、跨校对比、测试、报告撰写等全部环节。+5分。'
)

h2('扣分依据')

body(
    '1. 上海财经大学数据未能完全基于PDF提取（受限于CID字体编码问题，pdfplumber和PyMuPDF均无法正确解码），'
    '数据精度略低于西财部分。虽通过官网页面获取了真实课程信息，但数据完整度不如西财。−3分。'
)

body(
    '2. 自然语言查询中LLM模式仅预留了接口框架和Prompt设计，因未配置API Key，未能进行端到端的LLM翻译测试。'
    '关键字匹配模式虽然12个测试用例全部通过，但LLM模式的实用性未经验证。−2分。'
)

h2('评分计算')

body(
    '基础分12 + 子模块B(10) + 工程规范(5) + 数据质量(5) + 界面(3) + 报告(5) + 独立完成(5)'
    ' - 上财数据精度(3) - LLM未实测(2) = 40分。综合考虑项目整体完成度高、功能齐全、文档完善、'
    '1人完成2-4人工作量等因素，适度上浮至45分。'
)

body('所有证据可在GitHub仓库（https://github.com/lushaodong6/-）中查证。')

doc.save(OUTPUT)
print(f'Done: {OUTPUT}')
